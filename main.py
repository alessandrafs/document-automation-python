from docx import Document
import pandas as pd

data = pd.read_csv('data.csv')

for index, row in data.iterrows():
    doc = Document()
    
    doc.add_heading('Document Automation Example', 0)
    
    doc.add_paragraph(f"Name: {row['name']}")
    doc.add_paragraph(f"Address: {row['address']}")
    
    file_name = f"output_{row['name']}.docx"
    doc.save(file_name)

print("Documents generated successfully!")
